from langchain.docstore.document import Document
from docx import Document as DocxDocument
import logging

# Custom loader that uses DocsLoader internally and adds robust error handling.
class CustomDocsLoader:
    def __init__(self, file_path, original_filename=None):
        self.file_path = file_path
        self.original_filename = original_filename or file_path.split("/")[-1]

    def load(self):
        try:
            doc = DocxDocument(self.file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)

            return [Document(
                page_content=content,
                metadata={
                    "source": self.file_path,
                    "filename": self.original_filename
                }
            )]
        except Exception as e:
            logging.error(f"Error loading .docx file {self.file_path}: {e}")
            raise RuntimeError(f"Could not load .docx file {self.file_path}")